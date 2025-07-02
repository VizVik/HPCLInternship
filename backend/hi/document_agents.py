import os
import io
import base64
import pandas as pd
from dotenv import load_dotenv
from docx import Document as DocxDocument
import pdfplumber
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import google.generativeai as genai
from autogen import AssistantAgent

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

class DocumentAgent(AssistantAgent):
    def __init__(self):
        super().__init__(
            name="DocumentAgent",
            system_message="You are a helpful document assistant that answers questions based only on uploaded document content using RAG."
        )
        self.vectorstore = None
        self.embedding_model = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=os.getenv("GOOGLE_API_KEY")
        )
        self.model = genai.GenerativeModel("gemini-2.0-flash")

    def load_file(self, file_dict):
        filename = file_dict["name"]
        content = file_dict["content"]
        ext = os.path.splitext(filename)[1].lower()
        temp_path = os.path.join("temp", filename)
        os.makedirs("temp", exist_ok=True)

        # Decode base64 if needed
        if isinstance(content, str):
            try:
                content_bytes = base64.b64decode(content)
            except Exception:
                return content  # Already plain text or broken encoding
        else:
            content_bytes = content

        # Write to temp file
        with open(temp_path, "wb") as f:
            f.write(content_bytes)

        try:
            # === PDF Handling ===
            if ext == ".pdf":
                text = ""
                with pdfplumber.open(temp_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text if text.strip() else "❌ No readable text extracted from PDF."

            # === DOCX Handling ===
            elif ext == ".docx":
                try:
                    doc = DocxDocument(temp_path)
                    text_parts = []

                    # Paragraphs
                    text_parts.extend([para.text for para in doc.paragraphs if para.text.strip()])

                    # Tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_parts.append(cell.text)

                    full_text = "\n".join(text_parts)
                    return full_text if full_text.strip() else "❌ DOCX contains no readable text."

                except Exception as e:
                    return f"❌ Failed to read DOCX file: {e}"

            # === CSV Handling ===
            elif ext == ".csv":
                df = pd.read_csv(io.BytesIO(content_bytes), encoding="utf-8", errors="ignore")
                return df.to_string()

            # === XLSX Handling ===
            elif ext == ".xlsx":
                df = pd.read_excel(io.BytesIO(content_bytes), engine="openpyxl")
                return df.to_string()

            # === TXT Handling ===
            elif ext == ".txt":
                return content_bytes.decode("utf-8", errors="ignore")

            else:
                return "❌ Unsupported file format."

        except Exception as e:
            return f"❌ Error reading file: {e}"
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def create_vectorstore(self, text):
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(text)
        self.vectorstore = FAISS.from_texts(chunks, self.embedding_model)

    def get_context_from_query(self, query, k=4):
        if not self.vectorstore:
            return ""
        docs = self.vectorstore.similarity_search(query, k=k)
        return "\n\n".join(doc.page_content for doc in docs)

    def generate_answer(self, query, context):
        prompt = f"""
You are a helpful assistant. Use ONLY the context below to answer the question.

Context:
{context}

Question: {query}
Answer:
"""
        response = self.model.generate_content(prompt)
        return response.text.strip()

    async def run(self, input_data):
        if isinstance(input_data, dict) and "name" in input_data:
            content = self.load_file(input_data)
            if content.startswith("❌"):
                return content
            self.create_vectorstore(content)
            return "✅ Document indexed successfully. You can now ask questions."

        elif isinstance(input_data, str):
            context = self.get_context_from_query(input_data)
            if not context:
                return "❌ No document found. Please upload a document first."
            return self.generate_answer(input_data, context)

        else:
            return "❌ Invalid input to DocumentAgent."
